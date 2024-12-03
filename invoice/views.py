from django.core.mail import EmailMessage
from django.shortcuts import render
import os
from docx import Document
from django.http import FileResponse, JsonResponse
from django.views import View
from docx2pdf import convert
from django.conf import settings

def create_docx(data_dict, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    docx_file_path = os.path.join(output_dir, "output.docx")
    document = Document()

    for key, value in data_dict.items():
        document.add_heading(key, level=2)
        document.add_paragraph(str(value))

    document.save(docx_file_path)

    # Step 2: Convert to PDF
    pdf_file_path = os.path.join(output_dir, "output.pdf")
    convert(docx_file_path, pdf_file_path)

    return pdf_file_path


class Invoice(View):
    def get(self, request):
        data = {
                "Title": "Monthly Invoice",
                "Date": "December 1, 2024",
                "Author": "Jane Doe",
                "Introduction": (
                    "This report provides an overview of the sales performance "
                    "for the month of November 2024. Key metrics, trends, and "
                    "recommendations are included to help improve future performance."
                ),
                "Key Metrics": {
                    "Total Sales": "$50,000",
                    "New Customers": 150,
                    "Returning Customers": 120,
                    "Average Order Value": "$250",
                },
                "Summary": (
                    "Sales showed a 10% increase compared to October 2024. "
                    "The number of new customers has grown, indicating a successful marketing campaign."
                ),
                "Recommendations": [
                    "Increase the marketing budget for the holiday season.",
                    "Focus on customer retention strategies to boost repeat purchases.",
                    "Launch a loyalty program to reward returning customers."
                ],
            }
        output_dir = "/path/to/output/directory"

        pdf_path = create_docx(data, output_dir)

        if os.path.exists(pdf_path):
            email_subject = "Invoice"
            email_body = "This is an invoice reminder "
            email_to = "muhammadishaqskd224@gmail.com"
            email = EmailMessage(
                subject=email_subject,
                body=email_body,
                from_email=settings.EMAIL_HOST_USER,
                to=[email_to],  # User's email address
            )

            with open(pdf_path, 'rb') as pdf_file:
                email.attach("generated_report.pdf", pdf_file.read(), "application/pdf")

            try:
                email.send()
                contex = {
                    "email_to":email_to
                }
                return render(request, 'invoice.html', contex)
                # return JsonResponse({"message": "Email sent successfully!"}, status=200)
            except Exception as e:
                return JsonResponse({"error": str(e)}, status=500)
        else:
            return JsonResponse({"error": "Failed to generate PDF file."}, status=500)
